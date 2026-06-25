# check_docx_quality.py — validate docx template output quality
import sys
from docx import Document


def check_quality(docx_path):
    doc = Document(docx_path)
    issues = []

    # Check 1: Non-empty paragraphs
    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty_paras) == 0:
        issues.append("FAIL: 0 paragraphs with text — document appears blank")

    # Check 2: Table cell content
    total_cells = 0
    filled_cells = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                if cell.text.strip():
                    filled_cells += 1
    if total_cells > 0 and filled_cells / total_cells < 0.5:
        issues.append(
            f"FAIL: Only {filled_cells}/{total_cells} table cells contain text (< 50%)"
        )

    # Check 3: Font consistency
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and len(run.font.name) > 30:
                issues.append(
                    f"FAIL: Suspicious font name '{run.font.name[:50]}...'"
                )
                break

    # Summary
    if issues:
        print("QUALITY CHECK FAILED:")
        for issue in issues:
            print(f"  {issue}")
        return 1
    else:
        paras_with_text = len(non_empty_paras)
        table_count = len(doc.tables)
        total_cells_info = (
            f"{filled_cells}/{total_cells} cells filled"
            if total_cells > 0
            else "no tables"
        )
        print(
            f"QUALITY CHECK PASSED: {paras_with_text} paragraphs, "
            f"{table_count} tables ({total_cells_info})"
        )
        return 0


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python check_docx_quality.py <docx_path>")
        sys.exit(1)
    sys.exit(check_quality(sys.argv[1]))
