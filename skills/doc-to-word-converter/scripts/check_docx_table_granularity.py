#!/usr/bin/env python3
"""Check whether DOCX tables preserve row/cell granularity."""

from __future__ import annotations

import argparse
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def has_merged_cells(docx_path: Path) -> bool:
    with zipfile.ZipFile(docx_path, "r") as zf:
        for name in zf.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(zf.read(name))
            except ET.ParseError:
                continue
            if root.findall(".//w:gridSpan", NS) or root.findall(".//w:vMerge", NS):
                return True
    return False


def check_docx(docx_path: Path, min_rows: int, max_newlines_per_cell: int, no_merged_cells: bool) -> int:
    document = Document(str(docx_path))
    issues: list[str] = []

    if not document.tables:
        issues.append("FAIL: no tables found")
    else:
        largest_rows = max(len(table.rows) for table in document.tables)
        if largest_rows < min_rows:
            issues.append(f"FAIL: largest table has {largest_rows} rows; expected at least {min_rows}")

        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for col_index, cell in enumerate(row.cells, start=1):
                    newline_count = cell.text.count("\n")
                    if newline_count > max_newlines_per_cell:
                        issues.append(
                            "FAIL: table "
                            f"{table_index} row {row_index} col {col_index} has "
                            f"{newline_count} line breaks; expected <= {max_newlines_per_cell}"
                        )

    if no_merged_cells and has_merged_cells(docx_path):
        issues.append("FAIL: merged table cells found (gridSpan or vMerge)")

    if issues:
        print("TABLE GRANULARITY CHECK FAILED:")
        for issue in issues:
            print(f"  {issue}")
        return 1

    total_tables = len(document.tables)
    row_counts = [len(table.rows) for table in document.tables]
    print(f"TABLE GRANULARITY CHECK PASSED: {total_tables} tables, row counts={row_counts}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--min-rows", type=int, default=8)
    parser.add_argument("--max-newlines-per-cell", type=int, default=3)
    parser.add_argument("--no-merged-cells", action="store_true")
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"DOCX not found: {args.docx}", file=sys.stderr)
        return 2

    return check_docx(args.docx, args.min_rows, args.max_newlines_per_cell, args.no_merged_cells)


if __name__ == "__main__":
    raise SystemExit(main())
