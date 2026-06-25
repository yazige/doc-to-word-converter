#!/usr/bin/env python3
"""Create a simple Chinese contract DOCX from cleaned template text."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path


PAGE_BREAK_MARKERS = {"---PAGE BREAK---", "[PAGE BREAK]", "[分页]", "<分页>"}


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_normal_style(document) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(12)


def configure_page(document, orientation: str, margin_cm: float) -> None:
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT

    section = document.sections[0]
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)


def add_paragraph(document, text: str, *, title: bool = False, clause_heading: bool = False) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)

    if title:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, 16, True)
        paragraph_format.first_line_indent = None
        paragraph_format.space_after = Pt(6)
        return

    if clause_heading:
        run = paragraph.add_run(text)
        set_run_font(run, 12, True)
        paragraph_format.first_line_indent = Pt(24)
        return

    run = paragraph.add_run(text)
    set_run_font(run, 12, False)
    paragraph_format.first_line_indent = Pt(24)


def is_clause_heading(text: str) -> bool:
    return bool(re.match(r"^第[一二三四五六七八九十百]+条[ 　、:：]", text))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_text", type=Path, help="UTF-8 cleaned text file")
    parser.add_argument("output_docx", type=Path, help="DOCX file to create")
    parser.add_argument("--title-lines", type=int, default=1, help="Number of initial non-empty lines to center as title")
    parser.add_argument(
        "--orientation",
        choices=("portrait", "landscape"),
        default="portrait",
        help="Page orientation detected from the source document. Default is the safe fallback when detection fails.",
    )
    parser.add_argument(
        "--margin-cm",
        type=float,
        default=1.27,
        help="Page margin in centimeters. Default matches the template skill requirement.",
    )
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError:
        print("Missing dependency: python-docx. Install python-docx or use an environment that provides it.", file=sys.stderr)
        return 2

    if not args.input_text.exists():
        print(f"Input text not found: {args.input_text}", file=sys.stderr)
        return 2

    lines = args.input_text.read_text(encoding="utf-8").splitlines()
    document = Document()
    set_normal_style(document)
    configure_page(document, args.orientation, args.margin_cm)

    seen_title_lines = 0
    for raw_line in lines:
        line = raw_line.strip()
        if line in PAGE_BREAK_MARKERS:
            document.add_page_break()
            continue
        if not line:
            document.add_paragraph()
            continue
        if seen_title_lines < args.title_lines:
            add_paragraph(document, line, title=True)
            seen_title_lines += 1
            continue
        add_paragraph(document, line, clause_heading=is_clause_heading(line))

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.output_docx))
    print(args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
